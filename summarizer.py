from docx import Document
import openai
import os
from dotenv import load_dotenv
import fitz

OPENAI_API_KEY = 'sk-proj-xxxxxx'


# PyMuPDF

def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    return "\n".join([page.get_text() for page in doc])


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError("Unsupported file format. Use .docx or .pdf")


load_dotenv()
openai.api_key = OPENAI_API_KEY


def summarize_text(text):
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system",
             "content": "You are a legal assistant that summarizes legal documents into concise, readable briefs."},
            {"role": "user", "content": f"Summarize the following legal document:\n\n{text}"}
        ],
        temperature=0.5,
        max_tokens=1024
    )
    return response['choices'][0]['message']['content']


# def summarize_docx(file_path):
#     text = extract_text_from_docx(file_path)
#     if len(text) > 10000:
#         print("Warning: Document is very long. Consider chunking.")
#     summary = summarize_text(text)
#     return summary

def summarize_document(file_path):
    text = extract_text(file_path)
    if len(text) > 10000:
        print("Warning: Document is very long. Consider chunking.")
    return summarize_text(text)



if __name__ == "__main__":
    file_path = "Resume 25th Jan.pdf"#"your_legal_document.docx"
    summary = summarize_document(file_path)
    print("\n--- Summary ---\n")
    print(summary)
